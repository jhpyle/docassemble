import re
import os
import codecs
from copy import deepcopy
import tempfile
import string
from bs4 import BeautifulSoup, NavigableString, Tag
import docx
import qrcode
import qrcode.image.svg
from docxtpl import RichText
from ..config import daconfig
from ..functions import DALocalFile, roman, package_template_filename
from ..hooks import file_finder, fg_make_pdf_for_word_path
from ..logger import logmessage
from ..thread_context import this_thread
from .docx_subdoc import fix_subdoc
from .html import markdown_to_html
from .image_docx import image_for_docx
from .utils import (
    repeat_along,
    replace_fields,
    convert_pixels,
    get_default_image_width,
    convert_svg_to_eps,
    list_types,
    zerowidth,
    sanitize_xml,
)

NoneType = type(None)

def docx_filter(text, metadata=None, question=None):
    if metadata is None:
        metadata = {}
    text = text + "\n\n"
    text = re.sub(r'\[\[([^\]]*)\]\]', r'\1', text)
    text = re.sub(r'\[EMOJI ([^,\]]+), *([0-9A-Za-z.%]+)\]', lambda x: image_include_docx(x, question=question), text)
    text = re.sub(r'\[FILE ([^,\]]+), *([0-9A-Za-z.%]+), *([^\]]*)\]', lambda x: image_include_docx(x, question=question), text)
    text = re.sub(r'\[FILE ([^,\]]+), *([0-9A-Za-z.%]+)\]', lambda x: image_include_docx(x, question=question), text)
    text = re.sub(r'\[FILE ([^,\]]+)\]', lambda x: image_include_docx(x, question=question), text)
    text = re.sub(r'\[QR ([^,\]]+), *([0-9A-Za-z.%]+), *([^\]]*)\]', qr_include_docx, text)
    text = re.sub(r'\[QR ([^,\]]+), *([0-9A-Za-z.%]+)\]', qr_include_docx, text)
    text = re.sub(r'\[QR ([^\]]+)\]', qr_include_docx, text)
    text = re.sub(r'\[MAP ([^\]]+)\]', '', text)
    text = replace_fields(text)
    # text = re.sub(r'\[FIELD ([^\]]+)\]', '', text)
    text = re.sub(r'\[TARGET ([^\]]+)\]', '', text)
    text = re.sub(r'\[YOUTUBE[^ ]* ([^\]]+)\]', '', text)
    text = re.sub(r'\[VIMEO[^ ]* ([^\]]+)\]', '', text)
    text = re.sub(r'\\clearpage *\\clearpage', '', text)
    text = re.sub(r'\[START_INDENTATION\]', '', text)
    text = re.sub(r'\[STOP_INDENTATION\]', '', text)
    text = re.sub(r'\[BEGIN_CAPTION\](.+?)\[VERTICAL_LINE\]\s*(.+?)\[END_CAPTION\]', '', text, flags=re.DOTALL)
    text = re.sub(r'\[BEGIN_TWOCOL\](.+?)\[BREAK\]\s*(.+?)\[END_TWOCOL\]', '', text, flags=re.DOTALL)
    text = re.sub(r'\[TIGHTSPACING\] *', '', text)
    text = re.sub(r'\[SINGLESPACING\] *', '', text)
    text = re.sub(r'\[DOUBLESPACING\] *', '', text)
    text = re.sub(r'\[ONEANDAHALFSPACING\] *', '', text)
    text = re.sub(r'\[TRIPLESPACING\] *', '', text)
    text = re.sub(r'\[NBSP\]', ' ', text)
    text = re.sub(r'\[REDACTION_SPACE\]', "\u200B", text)
    text = re.sub(r'\[REDACTION_WORD ([^\]]+)\]', lambda x: repeat_along('█', x), text)
    text = re.sub(r'\[ENDASH\]', '--', text)
    text = re.sub(r'\[EMDASH\]', '---', text)
    text = re.sub(r'\[HYPHEN\]', '-', text)
    text = re.sub(r'\[CHECKBOX\]', '____', text)
    text = re.sub(r'\[BLANK\]', r'__________________', text)
    text = re.sub(r'\[BLANKFILL\]', r'__________________', text)
    text = re.sub(r'\[PAGEBREAK\] *', '', text)
    text = re.sub(r'\[PAGENUM\] *', '', text)
    text = re.sub(r'\[TOTALPAGES\] *', '', text)
    text = re.sub(r'\[SECTIONNUM\] *', '', text)
    text = re.sub(r'\[SKIPLINE\] *', '\n\n', text)
    text = re.sub(r'\[VERTICALSPACE\] *', '\n\n', text)
    text = re.sub(r'\[NEWLINE\] *', '\n\n', text)
    text = re.sub(r'\[NEWPAR\] *', '\n\n', text)
    text = re.sub(r'\[BR\] *', '\n\n', text)
    text = re.sub(r'\[TAB\] *', '', text)
    text = re.sub(r' *\[END\] *', r'\n', text)
    text = re.sub(r'\[BORDER\] *(.+?)\n *\n', r'\1\n\n', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[NOINDENT\] *(.+?)\n *\n', r'\1\n\n', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[FLUSHLEFT\] *(.+?)\n *\n', r'\1\n\n', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[FLUSHRIGHT\] *(.+?)\n *\n', r'\1\n\n', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[CENTER\] *(.+?)\n *\n', r'\1\n\n', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[BOLDCENTER\] *(.+?)\n *\n', r'**\1**\n\n', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[INDENTBY *([0-9\.]+ *[A-Za-z]+)\] *(.+?)\n *\n', r'\2', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[INDENTBY *([0-9\.]+ *[A-Za-z]+) *([0-9]+ *[A-Za-z]+)\] *(.+?)\n *\n', r'\3', text, flags=re.MULTILINE | re.DOTALL)
    return text


def docx_template_filter(text, question=None, replace_newlines=True):
    # logmessage('docx_template_filter')
    if text == 'True':
        return True
    if text == 'False':
        return False
    if text == 'None':
        return None
    text = re.sub(r'\[\[([^\]]*)\]\]', r'\1', text)
    text = re.sub(r'\[EMOJI ([^,\]]+), *([0-9A-Za-z.%]+)\]', lambda x: image_include_docx_template(x, question=question), text)
    text = re.sub(r'\[FILE ([^,\]]+), *([0-9A-Za-z.%]+), *([^\]]*)\]', lambda x: image_include_docx_template(x, question=question), text)
    text = re.sub(r'\[FILE ([^,\]]+), *([0-9A-Za-z.%]+)\]', lambda x: image_include_docx_template(x, question=question), text)
    text = re.sub(r'\[FILE ([^,\]]+)\]', lambda x: image_include_docx_template(x, question=question), text)
    text = re.sub(r'\[QR ([^,\]]+), *([0-9A-Za-z.%]+), *([^\]]*)\]', qr_include_docx_template, text)
    text = re.sub(r'\[QR ([^,\]]+), *([0-9A-Za-z.%]+)\]', qr_include_docx_template, text)
    text = re.sub(r'\[QR ([^\]]+)\]', qr_include_docx_template, text)
    text = re.sub(r'\[MAP ([^\]]+)\]', '', text)
    text = replace_fields(text)
    # text = re.sub(r'\[FIELD ([^\]]+)\]', '', text)
    text = re.sub(r'\[TARGET ([^\]]+)\]', '', text)
    text = re.sub(r'\[YOUTUBE[^ ]* ([^\]]+)\]', '', text)
    text = re.sub(r'\[VIMEO[^ ]* ([^\]]+)\]', '', text)
    text = re.sub(r'\\clearpage *\\clearpage', '', text)
    text = re.sub(r'\[START_INDENTATION\]', '', text)
    text = re.sub(r'\[STOP_INDENTATION\]', '', text)
    text = re.sub(r'\[BEGIN_CAPTION\](.+?)\[VERTICAL_LINE\]\s*(.+?)\[END_CAPTION\]', '', text, flags=re.DOTALL)
    text = re.sub(r'\[BEGIN_TWOCOL\](.+?)\[BREAK\]\s*(.+?)\[END_TWOCOL\]', '', text, flags=re.DOTALL)
    text = re.sub(r'\[TIGHTSPACING\] *', '', text)
    text = re.sub(r'\[SINGLESPACING\] *', '', text)
    text = re.sub(r'\[DOUBLESPACING\] *', '', text)
    text = re.sub(r'\[ONEANDAHALFSPACING\] *', '', text)
    text = re.sub(r'\[TRIPLESPACING\] *', '', text)
    text = re.sub(r'\[NBSP\]', ' ', text)
    text = re.sub(r'\[REDACTION_SPACE\]', "\u200B", text)
    # text = re.sub(r'\[REDACTION_SPACE\]', r'', text)
    text = re.sub(r'\[REDACTION_WORD ([^\]]+)\]', lambda x: repeat_along('█', x), text)
    # text = re.sub(r'\[REDACTION_WORD ([^\]]+)\]', lambda x: repeat_along('X', x), text)
    text = re.sub(r'\[ENDASH\]', '--', text)
    text = re.sub(r'\[EMDASH\]', '---', text)
    text = re.sub(r'\[HYPHEN\]', '-', text)
    text = re.sub(r'\\', '', text)
    text = re.sub(r'\[CHECKBOX\]', '____', text)
    text = re.sub(r'\[BLANK\]', r'__________________', text)
    text = re.sub(r'\[BLANKFILL\]', r'__________________', text)
    text = re.sub(r'\[PAGEBREAK\] *', '', text)
    text = re.sub(r'\[PAGENUM\] *', '', text)
    text = re.sub(r'\[TOTALPAGES\] *', '', text)
    text = re.sub(r'\[SECTIONNUM\] *', '', text)
    text = re.sub(r'\[SKIPLINE\] *', '</w:t><w:br/><w:t xml:space="preserve">', text)
    text = re.sub(r'\[VERTICALSPACE\] *', '</w:t><w:br/><w:br/><w:t xml:space="preserve">', text)
    text = re.sub(r'\[NEWLINE\] *', '</w:t><w:br/><w:t xml:space="preserve">', text)
    # text = re.sub(r'\n *\n', '[NEWPAR]', text)
    if replace_newlines:
        text = re.sub(r'\n', ' ', text)
    text = re.sub(r'\[NEWPAR\] *', '</w:t><w:br/><w:br/><w:t xml:space="preserve">', text)
    text = re.sub(r'\[TAB\] *', '\t', text)
    text = re.sub(r'\[NEWPAR\]', '</w:t><w:br/><w:br/><w:t xml:space="preserve">', text)
    text = re.sub(r' *\[END\] *', r'</w:t><w:br/><w:t xml:space="preserve">', text)
    text = re.sub(r'\[BORDER\] *', r'', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[NOINDENT\] *', r'', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[FLUSHLEFT\] *', r'', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[FLUSHRIGHT\] *', r'', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[CENTER\] *', r'', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[BOLDCENTER\] *', r'', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[INDENTBY *([0-9\.]+ *[A-Za-z]+)\] *(.+?)\n *\n', r'\2', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[INDENTBY *([0-9\.]+ *[A-Za-z]+) *([0-9]+ *[A-Za-z]+)\] *(.+?)\n *\n', r'\3', text, flags=re.MULTILINE | re.DOTALL)
    text = re.sub(r'\[BR\]', '</w:t><w:br/><w:t xml:space="preserve">', text)
    text = re.sub(r'\[SKIPLINE\]', '</w:t><w:br/><w:t xml:space="preserve">', text)
    text = re.sub(r'{([{%#])', '{' + zerowidth + r'\1', re.sub(r'([}%#])}', r'\1' + zerowidth + '}', text))
    return text


def image_include_docx(match, question=None):
    file_reference = match.group(1)
    if question and file_reference in question.interview.images:
        file_reference = question.interview.images[file_reference].get_reference()
    try:
        width = match.group(2)
        assert width != 'None'
        width = re.sub(r'^(.*)px', convert_pixels, width)
        if width == "full":
            width = '100%'
    except:
        width = get_default_image_width()
    if match.lastindex == 3:
        alt_text = match.group(3)
    else:
        alt_text = None
    if not alt_text:
        alt_text = ''
    file_info = file_finder(file_reference, question=question)
    if 'mimetype' in file_info and file_info['mimetype']:
        if re.search(r'^(audio|video)', file_info['mimetype']):
            return '[reference to file type that cannot be displayed]'
    if 'path' in file_info:
        if 'extension' in file_info:
            convert_svg_to_eps(file_info)
            if file_info['extension'] in ('docx', 'rtf', 'doc', 'odt'):
                if not os.path.isfile(file_info['path'] + '.pdf'):
                    fg_make_pdf_for_word_path(file_info['path'], file_info['extension'])
                output = '![' + alt_text + '](' + file_info['path'] + '.pdf){width=' + width + '}'
                return output
            if file_info['extension'] in ['png', 'jpg', 'gif', 'pdf', 'eps', 'jpe', 'jpeg']:
                output = '![' + alt_text + '](' + file_info['fullpath'] + '){width=' + width + '}'
                return output
    return '[invalid graphics reference]'


def image_include_docx_template(match, question=None):
    file_reference = match.group(1)
    if question and file_reference in question.interview.images:
        file_reference = question.interview.images[file_reference].get_reference()
    try:
        width = match.group(2)
        assert width != 'None'
        width = re.sub(r'^(.*)px', convert_pixels, width)
        if width == "full":
            width = '100%'
    except:
        width = get_default_image_width()
    if match.lastindex == 3:
        alt_text = match.group(3)
    else:
        alt_text = None
    file_info = file_finder(file_reference, question=question)
    if 'mimetype' in file_info and file_info['mimetype']:
        if re.search(r'^(audio|video)', file_info['mimetype']):
            return '[reference to file type that cannot be displayed]'
    if 'path' in file_info:
        convert_svg_to_eps(file_info)
        if 'mimetype' in file_info and file_info['mimetype']:
            if file_info['mimetype'] in ('text/markdown', 'text/plain'):
                with open(file_info['fullpath'], 'r', encoding='utf-8') as f:
                    contents = f.read()
                if file_info['mimetype'] == 'text/plain':
                    return contents
                return markdown_to_docx(contents, question, this_thread.misc.get('docx_template', None))
            if file_info['mimetype'] == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return str(include_docx_template(DALocalFile(file_info['fullpath'])))
            return str(image_for_docx(file_reference, question, this_thread.misc.get('docx_template', None), width=width, alt_text=alt_text))
    return '[reference to file that could not be found]'


def qr_include_docx(match):
    the_string = match.group(1)
    try:
        width = match.group(2)
        assert width != 'None'
        width = re.sub(r'^(.*)px', convert_pixels, width)
        if width == "full":
            width = '100%'
    except:
        width = get_default_image_width()
    if match.lastindex == 3:
        alt_text = match.group(3)
    else:
        alt_text = None
    if not alt_text:
        alt_text = ''
    im = qrcode.make(the_string)
    with tempfile.NamedTemporaryFile(prefix="datemp", suffix=".png", delete=False) as the_image:
        # this_thread.temporary_resources.add(the_image.name)
        im.save(the_image.name)
        output = '![' + alt_text + '](' + the_image.name + '){width=' + width + '}'
    return output


def qr_include_docx_template(match):
    the_string = match.group(1)
    try:
        width = match.group(2)
        assert width != 'None'
        width = re.sub(r'^(.*)px', convert_pixels, width)
        if width == "full":
            width = '100%'
    except:
        width = get_default_image_width()
    if match.lastindex == 3:
        alt_text = match.group(3)
    else:
        alt_text = None
    im = qrcode.make(the_string)
    with tempfile.NamedTemporaryFile(prefix="datemp", suffix=".png", delete=False) as the_image:
        im.save(the_image.name)
        return str(image_for_docx(DALocalFile(the_image.name), None, this_thread.misc.get('docx_template', None), width=width, alt_text=alt_text))


def Alpha(number):  # pylint: disable=invalid-name
    multiplier = int((number - 1) / 26)
    indexno = (number - 1) % 26
    return string.ascii_uppercase[indexno] * (multiplier + 1)


def alpha(number):
    multiplier = int((number - 1) / 26)
    indexno = (number - 1) % 26
    return string.ascii_lowercase[indexno] * (multiplier + 1)


def Roman_Numeral(number):  # pylint: disable=invalid-name
    return roman((number - 1) % 4000, case='upper')


def roman_numeral(number):
    return roman((number - 1) % 4000, case='lower')


class SoupParser:

    def __init__(self, tpl):
        self.paragraphs = [{'params': {'style': 'p', 'indentation': 0, 'list_number': 1}, 'runs': [RichText('')]}]
        self.current_paragraph = self.paragraphs[-1]
        self.run = self.current_paragraph['runs'][-1]
        self.bold = False
        self.center = False
        self.list_number = 1
        self.list_type = list_types[-1]
        self.italic = False
        self.underline = False
        self.strike = False
        self.indentation = 0
        self.style = 'p'
        self.still_new = True
        self.size = None
        self.charstyle = None
        self.color = None
        self.tpl = tpl

    def new_paragraph(self, classes, styles):
        if self.still_new:
            # logmessage("new_paragraph is still new and style is " + self.style + " and indentation is " + str(self.indentation))
            self.current_paragraph['params']['style'] = self.style
            self.current_paragraph['params']['indentation'] = self.indentation
            self.set_attribs(classes, styles)
            self.list_number += 1
            return
        # logmessage("new_paragraph where style is " + self.style + " and indentation is " + str(self.indentation))
        self.current_paragraph = {'params': {'style': self.style, 'indentation': self.indentation, 'list_number': self.list_number}, 'runs': [RichText('')]}
        self.set_attribs(classes, styles)
        self.list_number += 1
        self.paragraphs.append(self.current_paragraph)
        self.run = self.current_paragraph['runs'][-1]
        self.still_new = True

    def set_attribs(self, classes, styles):
        if 'dacenter' in classes:
            self.current_paragraph['params']['align'] = 'center'
        elif 'daflushright' in classes:
            self.current_paragraph['params']['align'] = 'end'
        else:
            self.current_paragraph['params']['align'] = 'start'
        if len(classes):
            if 'daspacingtight' in classes:
                self.current_paragraph['params']['spacing'] = 240
                self.current_paragraph['params']['after'] = 0
            elif 'daspacingsingle' in classes:
                self.current_paragraph['params']['spacing'] = 240
                self.current_paragraph['params']['after'] = 240
            elif 'daspacingdouble' in classes:
                self.current_paragraph['params']['spacing'] = 480
                self.current_paragraph['params']['after'] = 0
            elif 'daspacingoneandahalf' in classes:
                self.current_paragraph['params']['spacing'] = 260
                self.current_paragraph['params']['after'] = 0
            elif 'daspacingtriple' in classes:
                self.current_paragraph['params']['spacing'] = 700
                self.current_paragraph['params']['after'] = 0
        if styles:
            m = re.search(r'margin-left:([0-9\.]+)px', styles)
            if m:
                self.current_paragraph['params']['leftindent'] = 20 * int(m.group(1))
            m = re.search(r'margin-right:([0-9\.]+)px', styles)
            if m:
                self.current_paragraph['params']['rightindent'] = 20 * int(m.group(1))
            m = re.search(r'text-indent:([0-9\.]+)px', styles)
            if m:
                self.current_paragraph['params']['firstline'] = 20 * int(m.group(1))

    def __str__(self):
        output = ''
        for para in self.paragraphs:
            # logmessage("Got a paragraph where style is " + para['params'].get('style', 'undefined') + " and indentation is " + str(para['params'].get('indentation', 'undefined')))
            output += '<w:p><w:pPr><w:pStyle w:val="Normal"/>'
            if 'align' not in para['params']:
                para['params']['align'] = 'start'
            if para['params']['align'] == 'center':
                output += '<w:jc w:val="center"/>'
            elif para['params']['align'] == 'end':
                output += '<w:jc w:val="end"/>'
            if 'spacing' in para['params']:
                output += '<w:spacing w:before="0" w:after="' + str(para['params']['after']) + '" w:line="' + str(para['params']['spacing']) + '" w:lineRule="auto" w:beforeAutospacing="0" w:afterAutospacing="0"/>'
            if para['params']['style'] == 'ul' or para['params']['style'].startswith('ol'):
                if 'leftindent' in para['params']:
                    left_indent = para['params']['leftindent']
                else:
                    left_indent = 36*para['params']['indentation']
                if 'rightindent' in para['params']:
                    right_indent = para['params']['rightindent']
                else:
                    right_indent = 0
                output += '<w:ind w:left="' + str(left_indent) + '" w:right="' + str(right_indent) + '" w:hanging="0"/>'
            elif para['params']['style'] == 'blockquote':
                if 'spacing' not in para['params']:
                    output += '<w:spacing w:before="0" w:after="240" w:line="240" w:lineRule="auto" w:beforeAutospacing="0" w:afterAutospacing="0"/>'
                output += '<w:ind w:left="' + str(36*para['params']['indentation']) + '" w:right="' + str(36*para['params']['indentation']) + '" w:hanging="0"/>'
            elif 'leftindent' in para['params'] or 'rightindent' in para['params'] or 'firstline' in para['params']:
                if 'leftindent' in para['params']:
                    left_indent = para['params']['leftindent']
                else:
                    left_indent = 0
                if 'rightindent' in para['params']:
                    right_indent = para['params']['rightindent']
                else:
                    right_indent = 0
                if 'firstline' in para['params']:
                    first_line = para['params']['firstline']
                else:
                    first_line = 0
                output += '<w:ind w:left="' + str(left_indent) + '" w:right="' + str(right_indent) + '" w:hanging="0" w:firstLine="' + str(first_line) + '"/>'
            output += '<w:rPr></w:rPr></w:pPr>'
            if para['params']['style'] == 'ul':
                output += str(RichText("•\t"))
            if para['params']['style'] == 'ol1':
                output += str(RichText(str(para['params']['list_number']) + ".\t"))
            elif para['params']['style'] == 'olA':
                output += str(RichText(Alpha(para['params']['list_number']) + ".\t"))
            elif para['params']['style'] == 'ola':
                output += str(RichText(alpha(para['params']['list_number']) + ".\t"))
            elif para['params']['style'] == 'olI':
                output += str(RichText(Roman_Numeral(para['params']['list_number']) + ".\t"))
            elif para['params']['style'] == 'oli':
                output += str(RichText(roman_numeral(para['params']['list_number']) + ".\t"))
            for run in para['runs']:
                output += str(run)
            output += '</w:p>'
        return output

    def start_link(self, url):
        ref = self.tpl.docx._part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        self.current_paragraph['runs'].append('<w:hyperlink r:id="%s">' % (ref, ))
        self.new_run()
        self.still_new = False

    def end_link(self):
        self.current_paragraph['runs'].append('</w:hyperlink>')
        self.new_run()
        self.still_new = False

    def new_run(self):
        self.current_paragraph['runs'].append(RichText(''))
        self.run = self.current_paragraph['runs'][-1]

    def traverse(self, elem):
        for part in elem.contents:
            if isinstance(part, NavigableString):
                self.run.add(str(part), italic=self.italic, bold=self.bold, underline=self.underline, strike=self.strike, size=self.size, style=self.charstyle, color=self.color)
                self.still_new = False
            elif isinstance(part, Tag):
                # logmessage("Part name is " + str(part.name))
                if part.name == 'p':
                    if 'class' in part.attrs:
                        classes = part.attrs['class']
                    else:
                        classes = []
                    if 'style' in part.attrs:
                        styles = part.attrs['style']
                    else:
                        styles = ""
                    self.new_paragraph(classes, styles)
                    if 'dabold' in classes:
                        self.bold = True
                    self.traverse(part)
                    if 'dabold' in classes:
                        self.bold = False
                elif part.name == 'li':
                    if 'class' in part.attrs:
                        classes = part.attrs['class']
                    else:
                        classes = []
                    if 'style' in part.attrs:
                        styles = part.attrs['style']
                    else:
                        styles = ""
                    self.new_paragraph(classes, styles)
                    self.traverse(part)
                elif part.name == 'ul':
                    # logmessage("Entering a UL")
                    oldstyle = self.style
                    self.style = 'ul'
                    self.indentation += 10
                    self.traverse(part)
                    self.indentation -= 10
                    self.style = oldstyle
                    # logmessage("Leaving a UL")
                elif part.name == 'ol':
                    # logmessage("Entering a OL")
                    oldstyle = self.style
                    oldlistnumber = self.list_number
                    oldlisttype = self.list_type
                    if part.get('type', None) in list_types:
                        self.list_type = part['type']
                    else:
                        self.list_type = list_types[(list_types.index(self.list_type) + 1) % 5]
                    try:
                        self.list_number = int(part.get('start', 1))
                    except:
                        self.list_number = 1
                    self.style = 'ol' + self.list_type
                    self.indentation += 10
                    self.traverse(part)
                    self.indentation -= 10
                    self.list_type = oldlisttype
                    self.list_number = oldlistnumber
                    self.style = oldstyle
                    # logmessage("Leaving a OL")
                elif part.name == 'strong':
                    self.bold = True
                    self.traverse(part)
                    self.bold = False
                elif part.name == 'em':
                    self.italic = True
                    self.traverse(part)
                    self.italic = False
                elif part.name == 'strike':
                    self.strike = True
                    self.traverse(part)
                    self.strike = False
                elif part.name == 'u':
                    self.underline = True
                    self.traverse(part)
                    self.underline = False
                elif part.name == 'blockquote':
                    oldstyle = self.style
                    self.style = 'blockquote'
                    self.indentation += 20
                    self.traverse(part)
                    self.indentation -= 20
                    self.style = oldstyle
                elif re.match(r'h[1-6]', part.name):
                    oldsize = self.size
                    self.size = 60 - ((int(part.name[1]) - 1) * 10)
                    if 'class' in part.attrs:
                        classes = part.attrs['class']
                    else:
                        classes = []
                    if 'style' in part.attrs:
                        styles = part.attrs['style']
                    else:
                        styles = ""
                    self.new_paragraph(classes, styles)
                    self.bold = True
                    self.traverse(part)
                    self.bold = False
                    self.size = oldsize
                elif part.name == 'a':
                    self.start_link(part['href'])
                    if self.tpl.da_hyperlink_style:
                        self.charstyle = self.tpl.da_hyperlink_style
                    else:
                        self.underline = True
                        self.color = '#0000ff'
                    self.traverse(part)
                    if self.tpl.da_hyperlink_style:
                        self.charstyle = None
                    else:
                        self.underline = False
                        self.color = None
                    self.end_link()
                elif part.name == 'br':
                    self.run.add("\n", italic=self.italic, bold=self.bold, underline=self.underline, strike=self.strike, size=self.size, style=self.charstyle, color=self.color)
                    self.still_new = False
            else:
                logmessage("Encountered a " + part.__class__.__name__)


class InlineSoupParser:

    def __init__(self, tpl):
        self.runs = [RichText('')]
        self.run = self.runs[-1]
        self.bold = False
        self.italic = False
        self.underline = False
        self.indentation = 0
        self.style = 'p'
        self.strike = False
        self.size = None
        self.charstyle = None
        self.color = None
        self.tpl = tpl
        self.at_start = True
        self.list_number = 1
        self.list_type = list_types[-1]

    def new_paragraph(self):
        if self.at_start:
            self.at_start = False
        else:
            self.run.add("\n", italic=self.italic, bold=self.bold, underline=self.underline, strike=self.strike, size=self.size, style=self.charstyle, color=self.color)
        if self.indentation:
            self.run.add("\t" * self.indentation)
        if self.style == 'ul':
            self.run.add("•\t")
        if self.style == 'ol1':
            self.run.add(str(self.list_number) + ".\t")
            self.list_number += 1
        elif self.style == 'olA':
            self.run.add(Alpha(self.list_number) + ".\t")
            self.list_number += 1
        elif self.style == 'ola':
            self.run.add(alpha(self.list_number) + ".\t")
            self.list_number += 1
        elif self.style == 'olI':
            self.run.add(Roman_Numeral(self.list_number) + ".\t")
            self.list_number += 1
        elif self.style == 'oli':
            self.run.add(roman_numeral(self.list_number) + ".\t")
            self.list_number += 1
        # else:
        #     self.list_number = 1

    def __str__(self):
        output = ''
        for run in self.runs:
            output += str(run)
        return output

    def start_link(self, url):
        ref = self.tpl.docx._part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        self.runs.append('<w:hyperlink r:id="%s">' % (ref, ))
        self.new_run()

    def end_link(self):
        self.runs.append('</w:hyperlink>')
        self.new_run()

    def new_run(self):
        self.runs.append(RichText(''))
        self.run = self.runs[-1]

    def traverse(self, elem):
        for part in elem.contents:
            if isinstance(part, NavigableString):
                self.run.add(str(part), italic=self.italic, bold=self.bold, underline=self.underline, strike=self.strike, size=self.size, style=self.charstyle, color=self.color)
            elif isinstance(part, Tag):
                if part.name in ('p', 'blockquote'):
                    self.new_paragraph()
                    self.traverse(part)
                elif part.name == 'li':
                    self.new_paragraph()
                    self.traverse(part)
                elif part.name == 'ul':
                    oldstyle = self.style
                    self.style = 'ul'
                    self.indentation += 1
                    self.traverse(part)
                    self.indentation -= 1
                    self.style = oldstyle
                elif part.name == 'ol':
                    oldstyle = self.style
                    oldlistnumber = self.list_number
                    oldlisttype = self.list_type
                    if part.get('type', None) in list_types:
                        self.list_type = part['type']
                    else:
                        self.list_type = list_types[(list_types.index(self.list_type) + 1) % 5]
                    try:
                        self.list_number = int(part.get('start', 1))
                    except:
                        self.list_number = 1
                    self.style = 'ol' + self.list_type
                    self.indentation += 1
                    self.traverse(part)
                    self.indentation -= 1
                    self.list_type = oldlisttype
                    self.list_number = oldlistnumber
                    self.style = oldstyle
                elif part.name == 'strong':
                    self.bold = True
                    self.traverse(part)
                    self.bold = False
                elif part.name == 'em':
                    self.italic = True
                    self.traverse(part)
                    self.italic = False
                elif part.name == 'strike':
                    self.strike = True
                    self.traverse(part)
                    self.strike = False
                elif part.name == 'u':
                    self.underline = True
                    self.traverse(part)
                    self.underline = False
                elif re.match(r'h[1-6]', part.name):
                    oldsize = self.size
                    self.size = 60 - ((int(part.name[1]) - 1) * 10)
                    self.bold = True
                    self.traverse(part)
                    self.bold = False
                    self.size = oldsize
                elif part.name == 'a':
                    self.start_link(part['href'])
                    if self.tpl.da_hyperlink_style:
                        self.charstyle = self.tpl.da_hyperlink_style
                    else:
                        self.underline = True
                        self.color = '#0000ff'
                    self.traverse(part)
                    if self.tpl.da_hyperlink_style:
                        self.charstyle = None
                    else:
                        self.underline = False
                        self.color = None
                    self.end_link()
                elif part.name == 'br':
                    self.run.add("\n", italic=self.italic, bold=self.bold, underline=self.underline, strike=self.strike, size=self.size, style=self.charstyle, color=self.color)
            else:
                logmessage("Encountered a " + part.__class__.__name__)


def inline_markdown_to_docx(text, question, tpl):
    old_context = this_thread.evaluation_context
    this_thread.evaluation_context = None
    try:
        text = str(text)
    except:
        this_thread.evaluation_context = old_context
        raise
    this_thread.evaluation_context = old_context
    source_code = markdown_to_html(text, do_terms=False)
    source_code = re.sub(r"\n", ' ', source_code)
    source_code = re.sub(r">\s+<", '><', source_code)
    soup = BeautifulSoup('<html>' + source_code + '</html>', 'html.parser')
    parser = InlineSoupParser(tpl)
    for elem in soup.find_all(recursive=False):
        parser.traverse(elem)
    output = str(parser)
    return docx_template_filter(output, question=question, replace_newlines=False)


def markdown_to_docx(text, question, tpl):
    old_context = this_thread.evaluation_context
    this_thread.evaluation_context = None
    try:
        text = str(text)
    except:
        this_thread.evaluation_context = old_context
        raise
    this_thread.evaluation_context = old_context
    if daconfig.get('new markdown to docx', False):
        source_code = markdown_to_html(text, do_terms=False)
        source_code = re.sub(r"\n", ' ', source_code)
        source_code = re.sub(r">\s+<", '><', source_code)
        soup = BeautifulSoup('<html>' + source_code + '</html>', 'html.parser')
        parser = SoupParser(tpl)
        for elem in soup.find_all(recursive=False):
            parser.traverse(elem)
        output = str(parser)
        # logmessage(output)
        return docx_template_filter(output, question=question)
    return inline_markdown_to_docx(text, question, tpl)


def include_docx_template(template_file, **kwargs):
    """Include the contents of one docx file inside another docx file."""
    use_jinja = kwargs.pop('_use_jinja2', True)
    if this_thread.evaluation_context is None:
        return 'ERROR: not in a docx file'
    if template_file.__class__.__name__ in ('DAFile', 'DAFileList', 'DAFileCollection', 'DALocalFile', 'DAStaticFile'):
        template_path = template_file.path()
    else:
        template_path = package_template_filename(template_file, package=this_thread.current_package)
    sd = this_thread.misc['docx_template'].new_subdoc()
    sd.subdocx = docx.Document(template_path)
    change_numbering = bool(kwargs.pop('change_numbering', True))
    if '_inline' in kwargs:
        single_paragraph = True
        del kwargs['_inline']
    else:
        single_paragraph = False

    # We need to keep a copy of the subdocs so we can fix up the master template in the end (in parse.py)
    # Given we're half way through processing the template, we can't fix the master template here
    # we have to do it in post
    if 'docx_subdocs' not in this_thread.misc:
        this_thread.misc['docx_subdocs'] = []
    this_thread.misc['docx_subdocs'].append({'subdoc': deepcopy(sd.subdocx), 'change_numbering': change_numbering})

    # Fix the subdocs before they are included in the template
    fix_subdoc(this_thread.misc['docx_template'], {'subdoc': sd.subdocx, 'change_numbering': change_numbering})

    first_paragraph = sd.subdocx.paragraphs[0]

    if not use_jinja:
        if single_paragraph:
            return re.sub(r'<w:p[^>]*>\s*(.*)</w:p>\s*', r'\1', sanitize_xml(str(first_paragraph._p.xml)), flags=re.DOTALL)
        return sanitize_xml(str(sd))

    for key, val in kwargs.items():
        if hasattr(val, 'instanceName'):
            the_repr = val.instanceName
        elif isinstance(val, (int, float, bool, NoneType)):
            the_repr = val
        else:
            the_repr = '_codecs.decode(_array.array("b", "' + re.sub(r'\n', '', codecs.encode(bytearray(val, encoding='utf-8'), 'base64').decode()) + '".encode()), "base64").decode()'
        first_paragraph.insert_paragraph_before(str("{%%p set %s = %s %%}" % (key, the_repr)))
    if 'docx_include_count' not in this_thread.misc:
        this_thread.misc['docx_include_count'] = 0
    this_thread.misc['docx_include_count'] += 1
    if single_paragraph:
        return re.sub(r'<w:p[^>]*>\s*(.*)</w:p>\s*', r'\1', str(first_paragraph._p.xml), flags=re.DOTALL)
    return sd
